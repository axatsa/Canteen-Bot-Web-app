import os
import pytest
from backend.app.export import fill_docx_template, TEMPLATES_DIR, EXPORTS_DIR
from docx import Document

def test_export_content_validation():
    # Setup test data
    template_path = os.path.join(TEMPLATES_DIR, 'land_template.docx')
    if not os.path.exists(template_path):
        pytest.skip("Template land_template.docx not found")

    context = {
        'order_id': 'test-123',
        'branch': 'Белтепа-Land',
        'day': '21',
        'month_name': 'Апреля',
        'year': '2026',
        'snabjenec_name': 'Иван Снабженец',
        'supplier_name': 'ООО Поставка',
        'all_items': [
            {
                'product_name': 'Картофель',
                'unit': 'кг',
                'ordered_qty': 50,
                'received_qty': 45
            },
            {
                'product_name': 'Лук',
                'unit': 'кг',
                'ordered_qty': 10,
                'received_qty': 10
            }
        ]
    }

    # Generate document
    out_path = fill_docx_template(template_path, context)
    assert out_path is not None
    assert os.path.exists(out_path)

    try:
        # Verify document content
        doc = Document(out_path)
        
        # Check if branch name is in any paragraph
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Белтепа-Land" in all_text
        assert "Иван Снабженец" in all_text
        
        # Check table content
        table_found = False
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join([cell.text for cell in row.cells])
                if "Картофель" in row_text:
                    table_found = True
                    assert "50" in row_text
                    assert "45" in row_text
        
        assert table_found, "Product 'Картофель' not found in document tables"

    finally:
        # Clean up
        if os.path.exists(out_path):
            os.remove(out_path)

def test_export_missing_items():
    template_path = os.path.join(TEMPLATES_DIR, 'land_template.docx')
    if not os.path.exists(template_path):
        pytest.skip("Template land_template.docx not found")

    context = {
        'order_id': 'test-missing',
        'branch': 'Учтепа-Land',
        'all_items': []
    }
    
    out_path = fill_docx_template(template_path, context)
    assert out_path is not None
    os.remove(out_path)
