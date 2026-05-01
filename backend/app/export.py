import os
import re
import uuid
import logging
from typing import Optional
from datetime import datetime

logger = logging.getLogger(__name__)

TEMPLATES_DIR = os.getenv('TEMPLATES_DIR', os.path.join(os.path.dirname(__file__), '..', 'templates'))
EXPORTS_DIR = os.getenv('EXPORTS_DIR', os.path.join(os.path.dirname(__file__), '..', 'exports'))


def ensure_dirs():
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    os.makedirs(EXPORTS_DIR, exist_ok=True)


def _set_cell_text(cell, text: str):
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = str(text)
        for r in para.runs[1:]:
            r.text = ''
    else:
        from docx.oxml import OxmlElement
        r_el = OxmlElement('w:r')
        t_el = OxmlElement('w:t')
        t_el.text = str(text)
        r_el.append(t_el)
        para._p.append(r_el)


def _normalize_product_name(raw: str) -> str:
    if not raw: return ''
    name = raw.split('(')[0].strip()
    name = ''.join(re.findall(r'[a-zA-Zа-яА-Я0-9ёЁ]+', name.lower()))
    return name


def fill_docx_template(template_path: str, context: dict) -> Optional[str]:
    try:
        from docx import Document
    except ImportError:
        logger.error('python-docx not installed')
        return None

    ensure_dirs()
    try:
        doc = Document(template_path)

        day         = context.get('day', '')
        month_name  = context.get('month_name', '')
        time_str    = context.get('time_str', '')
        branch      = context.get('branch', '')
        snabjenec   = context.get('snabjenec_name', '')
        supplier    = context.get('supplier_name', '')
        recipient   = context.get('recipient_name', '') or snabjenec

        for para in doc.paragraphs:
            full = para.text
            if 'Дата:' in full:
                counter = [0]
                def _replace_underscores(m, _month=month_name, _time=time_str, _c=counter):
                    _c[0] += 1
                    return f' {_month} ' if _c[0] == 1 else f' {_time} '
                
                for r in para.runs:
                    r.text = re.sub(r'«[^»]*»', f'« {day} »', r.text)
                    r.text = re.sub(r'_{5,}', _replace_underscores, r.text)
            
            if 'Филиал:' in full:
                for r in para.runs:
                    if '_' in r.text:
                        r.text = re.sub(r'_{5,}', branch, r.text)
                if '____' in para.text and branch not in para.text:
                    para.text = para.text.replace('________________', branch).replace('________', branch)
            
            if full.strip().startswith(('Cнабженец:', 'Поставщик:', 'Получатель:')):
                for r in para.runs:
                    r.text = re.sub(r'_{5,}', (snabjenec if 'Cнабженец' in full else supplier if 'Поставщик' in full else recipient), r.text)

        if doc.tables:
            table = doc.tables[0]
            all_items = context.get('all_items', [])
            extra_items = context.get('extra_items', [])
            order_lookup = { _normalize_product_name(item['product_name']): item for item in all_items }
            matched_keys = set()
            empty_rows = []

            for row in table.rows:
                cells = row.cells
                ucells = []
                last_id = None
                for c in cells:
                    if id(c) != last_id:
                        ucells.append(c)
                        last_id = id(c)

                if len(ucells) < 2: continue
                num_text = ucells[0].text.strip()
                name_text = ucells[1].text.strip()

                if num_text.isdigit() and not name_text:
                    empty_rows.append((row, ucells))
                    continue

                if not num_text.isdigit() or not name_text:
                    continue

                norm = _normalize_product_name(name_text)
                matched = None
                if norm in order_lookup and norm not in matched_keys:
                    matched = order_lookup[norm]
                    matched_keys.add(norm)
                else:
                    for ok, ov in order_lookup.items():
                        if ok in matched_keys: continue
                        if ok in norm or norm in ok:
                            matched = ov
                            matched_keys.add(ok)
                            break

                if matched:
                    if len(ucells) >= 6:
                        _set_cell_text(ucells[2], str(matched.get('unit', '')))
                        _set_cell_text(ucells[3], str(matched.get('ordered_qty', '')))
                        _set_cell_text(ucells[4], str(matched.get('price', '')))
                        _set_cell_text(ucells[5], str(matched.get('sum', '')))
                    elif len(ucells) >= 5:
                        _set_cell_text(ucells[2], str(matched.get('unit', '')))
                        _set_cell_text(ucells[3], str(matched.get('ordered_qty', '')))
                        _set_cell_text(ucells[4], str(matched.get('received_qty', '')))
                    elif len(ucells) >= 4:
                        _set_cell_text(ucells[2], str(matched.get('ordered_qty', '')))
                        _set_cell_text(ucells[3], str(matched.get('received_qty', '')))

            still_unmatched = [ov for ok, ov in order_lookup.items() if ok not in matched_keys] + list(extra_items)
            for (row, ucells), item in zip(empty_rows, still_unmatched):
                _set_cell_text(ucells[1], item['product_name'])
                if len(ucells) >= 6:
                    _set_cell_text(ucells[2], str(item.get('unit', '')))
                    _set_cell_text(ucells[3], str(item.get('ordered_qty', '')))
                    _set_cell_text(ucells[4], str(item.get('price', '')))
                    _set_cell_text(ucells[5], str(item.get('sum', '')))
                elif len(ucells) >= 5:
                    _set_cell_text(ucells[2], str(item.get('unit', '')))
                    _set_cell_text(ucells[3], str(item.get('ordered_qty', '')))
                    _set_cell_text(ucells[4], str(item.get('received_qty', '')))
                else:
                    _set_cell_text(ucells[2], str(item.get('ordered_qty', '')))
                    _set_cell_text(ucells[3], str(item.get('received_qty', '')))

        out_name = f'order_{context.get("order_id", uuid.uuid4())}_{uuid.uuid4().hex[:8]}.docx'
        out_path = os.path.join(EXPORTS_DIR, out_name)
        doc.save(out_path)
        return out_path
    except Exception as e:
        logger.exception(f'Error filling DOCX template: {e}')
        return None

BRANCH_NAMES = {
    'beltepa_land':          'Белтепа-Land',
    'uchtepa_land':          'Учтепа-Land',
    'rakat_land':            'Ракат-Land',
    'mukumiy_land':          'Мукумий-Land',
    'yunusabad_land':        'Юнусабад-Land',
    'novoi_land':            'Новои-Land',
    'novza_school':          'Новза-School',
    'uchtepa_school':        'Учтепа-School',
    'almazar_school':        'Алмазар-School',
    'general_uzakov_school': 'Генерал Узаков-School',
    'namangan_school':       'Наманган-School',
    'novoi_school':          'Новои-School',
}

MONTH_NAMES = {
    1: 'Января', 2: 'Февраля', 3: 'Марта', 4: 'Апреля',
    5: 'Мая', 6: 'Июня', 7: 'Июля', 8: 'Августа',
    9: 'Сентября', 10: 'Октября', 11: 'Ноября', 12: 'Декабря',
}

def build_export_context(order_details: dict, names: dict = None) -> dict:
    order           = order_details.get('order', {})
    delivery        = order_details.get('delivery', {})
    delivered       = order_details.get('delivered_items', [])
    not_delivered   = order_details.get('not_delivered_items', [])
    extra           = order_details.get('extra_items', [])
    
    source_items = delivered + not_delivered
    if not source_items:
        source_items = [
            {
                'product_name': p.get('product_name', '') or p.get('name', ''),
                'unit': p.get('unit', ''),
                'ordered_qty': p.get('ordered_qty', 0) or p.get('quantity', 0),
                'received_qty': '',
                'status': 'pending'
            }
            for p in order_details.get('ordered_products', [])
            if (p.get('ordered_qty', 0) or p.get('quantity', 0)) > 0
        ]

    created_at_str = order.get('createdAt', '')
    try:
        dt = datetime.fromisoformat(created_at_str.replace('Z', '+00:00'))
    except:
        dt = datetime.now()

    sent_at_str = order.get('sent_to_financier_at', '')
    try:
        sent_dt = datetime.fromisoformat(sent_at_str) if sent_at_str else dt
    except:
        sent_dt = dt
    time_str = sent_dt.strftime('%H:%M')

    day         = str(dt.day)
    month       = MONTH_NAMES.get(dt.month, str(dt.month))
    year        = str(dt.year)
    total_ordered  = sum(i.get('ordered_qty', 0) for i in source_items)
    total_received = sum(i.get('received_qty', 0) if isinstance(i.get('received_qty'), (int, float)) else 0 for i in source_items)

    all_items = []
    for idx, item in enumerate(source_items, start=1):
        price = item.get('price') or ''
        qty = item.get('ordered_qty') or 0
        all_items.append({
            'number':       idx,
            'product_name': item.get('product_name', ''),
            'unit':         item.get('unit', ''),
            'ordered_qty':  item.get('ordered_qty', 0) or '',
            'received_qty': item.get('received_qty', 0) if item.get('received_qty') != '' else '',
            'price':        price,
            'sum':          round(price * qty, 2) if price != '' else '',
            'status':       item.get('status', ''),
        })

    snabjenec_name = (names or {}).get('snabjenec_name', '')
    supplier_name  = (names or {}).get('supplier_name', '')

    return {
        'order_id':         order.get('id', ''),
        'branch':           BRANCH_NAMES.get(order.get('branch', ''), order.get('branch', '')),
        'date':             dt.strftime('%d.%m.%Y'),
        'day':              day,
        'month_name':       month,
        'year':             year,
        'time_str':         time_str,
        'delivered_items':  delivered,
        'not_delivered_items': not_delivered,
        'extra_items':      extra,
        'all_items':        all_items,
        'total_ordered':    total_ordered,
        'total_received':   total_received,
        'completion_rate':  delivery.get('completion_rate', ''),
        'snabjenec_name':   snabjenec_name,
        'supplier_name':    supplier_name,
        'chef_name':        (names or {}).get('chef_name', ''),
        'recipient_name':   snabjenec_name,
        'sender_name':      supplier_name,
    }
