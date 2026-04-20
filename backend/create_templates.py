"""
Generate SCHOOL and LAND .docx templates with correct docxtpl tags.
Run inside the api container:
    docker compose exec api python create_templates.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), 'templates')
os.makedirs(TEMPLATES_DIR, exist_ok=True)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ['top', 'left', 'bottom', 'right']:
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:color'), '000000')
        tcPr.append(tag)


def set_bold(para, text, size=11, align=None):
    para.clear()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if align:
        para.alignment = align


def add_cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    para.alignment = align
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def create_template(filename, branch_label):
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Акт приема-передачи Продуктов для Кухни № ')
    run.bold = True
    run.font.size = Pt(13)
    run2 = title.add_run('{{ order_id }}')
    run2.bold = True
    run2.font.size = Pt(13)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run('Дата: «').bold = False
    r = date_para.add_run(' {{ day }} ')
    r.bold = True
    date_para.add_run('» ')
    r2 = date_para.add_run('{{ month_name }}, {{ year }}')
    r2.bold = True
    date_para.add_run(' г    Время: ')
    r3 = date_para.add_run('{{ time }}')
    r3.bold = True

    # Branch
    branch_para = doc.add_paragraph()
    branch_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    b = branch_para.add_run(f'{branch_label} Филиал:  ')
    b.bold = True
    b.font.size = Pt(12)
    r4 = branch_para.add_run('{{ branch }}')
    r4.bold = True
    r4.font.size = Pt(12)

    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'

    # Set column widths
    widths = [Cm(1.2), Cm(7.5), Cm(2.5), Cm(3.0), Cm(3.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    # Header row
    headers = ['No', 'Наименование товара', 'Ед. изм.', 'Кол-во Сдал', 'Кол-во Принял']
    for i, hdr in enumerate(headers):
        add_cell_text(table.rows[0].cells[i], hdr, bold=True, size=10)
        set_cell_border(table.rows[0].cells[i])

    # Template row — loop tags must be in their OWN paragraph (docxtpl requirement)
    row = table.rows[1]

    # Cell 0: loop START tag in separate paragraph, then number
    c0 = row.cells[0]
    c0.paragraphs[0].clear()
    c0.paragraphs[0].add_run('{%tr for item in all_items %}').font.size = Pt(1)
    p = c0.add_paragraph('{{ item.number }}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    set_cell_border(c0)

    # Cell 1: product name
    c1 = row.cells[1]
    add_cell_text(c1, '{{ item.product_name }}', size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_border(c1)

    # Cell 2: unit
    c2 = row.cells[2]
    add_cell_text(c2, '{{ item.unit }}', size=10)
    set_cell_border(c2)

    # Cell 3: ordered qty
    c3 = row.cells[3]
    add_cell_text(c3, '{{ item.ordered_qty }}', size=10)
    set_cell_border(c3)

    # Cell 4: received qty, then loop END tag in separate paragraph
    c4 = row.cells[4]
    c4.paragraphs[0].clear()
    c4.paragraphs[0].add_run('{{ item.received_qty }}').font.size = Pt(10)
    c4.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4 = c4.add_paragraph('{%tr endfor %}')
    p4.runs[0].font.size = Pt(1)
    set_cell_border(c4)

    doc.add_paragraph()
    doc.add_paragraph()

    # Signatures
    sigs = [
        ('Снабженец:', '{{ snabjenec_name }}'),
        ('Поставщик:', '{{ supplier_name }}'),
        ('Получатель:', '{{ chef_name }}'),
    ]
    for label, var in sigs:
        p = doc.add_paragraph()
        r_label = p.add_run(f'{label}  ')
        r_label.bold = True
        r_label.font.size = Pt(11)
        r_var = p.add_run(var)
        r_var.bold = True
        r_var.font.size = Pt(11)
        p.add_run('          подпись: _______________   время: _______')

    doc.add_paragraph()

    # Footer note
    note = doc.add_paragraph('Товар получен в полном объёме, без видимых повреждений.')
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].font.size = Pt(11)

    path = os.path.join(TEMPLATES_DIR, filename)
    doc.save(path)
    print(f'✅ Создан: {path}')
    return path


if __name__ == '__main__':
    create_template('school_template.docx', 'SCHOOL')
    create_template('land_template.docx', 'LAND')
    print('\nДобавь шаблоны в БД через вкладку Настройки или запусти register_templates.py')
